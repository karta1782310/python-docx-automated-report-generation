#!/usr/bin/env python
# coding: utf-8

# In[1]:

import os
import copy
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm

from math import ceil

from openpyxl import load_workbook

from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import Cm
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 處理字串對齊
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT  # 處理表格對齊
from docx.enum.table import WD_ALIGN_VERTICAL

from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement

GRAY = RGBColor(204, 204, 204)
BLACK = RGBColor(0, 0, 0)


# In[2]:


class Table():

    def set_cell_color(cell, rgbColor):
        shading_elm_1 = parse_xml(
            r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value=rgbColor))
        cell._tc.get_or_add_tcPr().append(shading_elm_1)

    def set_content_font(table, size=Pt(14), bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
        for row in table.rows:
            Table.set_row_height(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                Parag.set_font(cell.paragraphs, size, bold, align)

    def col_widths(table, *width):
        for i, col in enumerate(table.columns):
            if i < len(width):
                col.width = Cm(width[i])

    def row_height(table, height):
        for i, row in enumerate(table.rows):
            row.height = height

    # https://stackoverflow.com/questions/37532283/python-how-to-adjust-row-height-of-table-in-docx
    def set_row_height(row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), "100")
        trHeight.set(qn('w:hRule'), "auto")
        trPr.append(trHeight)

    # https://stackoverflow.com/questions/55545494/in-python-docx-how-do-i-delete-a-table-row
    def remove_row(table, row):
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)

    def add(table, data):
        for row in data:
            row_cell = table.add_row().cells
            for i, cell in enumerate(row):
                row_cell[i].text = str(cell)


# In[3]:


class Parag():

    def set_font(parags, size=Pt(14), bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):

        if isinstance(parags, list):
            for paragraph in parags:
                paragraph.paragraph_format.alignment = align
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for run in paragraph.runs:
                    run.bold = bold
                    run.font.size = size
                    run.font.name = 'Arial'  # 設置英文字體
                    run._element.rPr.rFonts.set(
                        qn('w:eastAsia'), u'標楷體')  # 設置中文字體

        else:
            parags.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            parags.paragraph_format.alignment = align
            for run in parags.runs:
                run.bold = bold
                run.font.size = size
                run.font.name = 'Arial'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

    def delete(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    # Insert a new paragraph after the given paragraph.
    def insert_paragraph_after(paragraph, text=None, style=None):
        new_p = OxmlElement("w:p")
        paragraph._p.addnext(new_p)
        new_para = Paragraph(new_p, paragraph._parent)
        if text:
            new_para.add_run(text)
        if style is not None:
            new_para.style = style
        return new_para


# In[4]:


def prevent_document_break(document):
    tags = document.element.xpath('//w:tr/w:tc/w:tbl/w:tr/w:tc/w:p')
    rows = len(tags)
    for row in range(0, rows):
        tag = tags[row]  # Specify which <w:r> tag you want
        child = OxmlElement('w:cantSplit')  # Create arbitrary tag
        tag.append(child)  # Append in the new tag


# In[5]:


def plt_bar(title, names, values, picture):
    plt.style.use('seaborn-whitegrid')

    plt.figure(figsize=(9, 4))
    plt.title(title, size=22)
    plt.ylim(0, (max(values) * 6/5 + 20) // 10 * 10)

    # for chinese show
    # font_path = '/Users/SystexB/Library/Fonts/Kaiu.ttf'
    # if os.path.isfile(font_path):
    #     font = fm.FontProperties(fname=font_path,size=16)
    #     plt.xticks(fontproperties=font)
    #     plt.title("風險", fontproperties=font)
    # else:
    plt.rcParams['font.sans-serif'] = ['DFKai-SB']  # 替換sans-serif字型
    plt.rcParams['axes.unicode_minus'] = False  # 解決座標軸負數的負號顯示問題
    plt.rcParams.update({'font.size': 16})
    plt.xticks(size=22)

    bar = plt.bar(names, values, color=[
                  'indianred', 'cornflowerblue', 'yellowgreen'])
    for rect in bar:
        h = rect.get_height()
        plt.text(rect.get_x()+rect.get_width()/2.0, h, "%d\n" %
                 h, ha='center', va='bottom')

    plt.savefig(picture, dpi=300)
    # plt.show()


# In[6]:


def risk_bar(doc, wb, sheet, ips=[]):

    names = ['高風險', '中風險', '低風險']
    values = [0, 0, 0]
    level = ""
    tmp_ips = copy.deepcopy(ips)

    # calculate needed data
    for row in wb[sheet].values:
        if tmp_ips:
            for ip in tmp_ips:
                if row[0] == ip:
                    level = row[7]
        else:
            level = row[7]

        if level == '高':
            values[0] += 1
        elif level == "中":
            values[1] += 1
        elif level == "低":
            values[2] += 1

        level = ""

    # find the mark
    for table in doc.tables:
        if "riskBar" in table.cell(0, 0).paragraphs[0].text:
            table.allow_autofit = True

            # add pic in table
            plt_bar('風險', names, values, "risk_bar.png")

            main_cell = table.cell(0, 0)
            main_cell.paragraphs[0].text = ""

            parag = main_cell.paragraphs[0]
            parag.add_run().add_picture("risk_bar.png", width=Cm(15))
            parag.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # title = main_cell.add_paragraph()
            # title.text = ("圖: 主機弱點掃描結果風險數量分布圖")
            # Parag.set_font(title)

    os.remove("risk_bar.png")

    # 替換標定資料
    for parag in doc.paragraphs:
        row = parag.text
        if "weaknessCount" in row or "highRiskWeaknessCount" in row or "middleRiskWeaknessCount" in row or "lowRiskWeaknessCount" in row:
            row = row.replace("weaknessCount", str(
                values[0]+values[1]+values[2]), 1)
            row = row.replace("highRiskWeaknessCount", str(values[0]), 1)
            row = row.replace("middleRiskWeaknessCount", str(values[1]), 1)
            row = row.replace("lowRiskWeaknessCount", str(values[2]), 1)
            parag.text = row


# In[7]:


def ip_list(doc, wb, sheet, ips=[]):

    tmp_ips = copy.deepcopy(ips)

    # calculate needed data
    if not tmp_ips:
        for row in wb[sheet].values:
            if row[0] not in tmp_ips and row[0]:
                tmp_ips.append(row[0])

        del(tmp_ips[0])  # delete "IP初始"

        # sort ip
        tmp_ips = sorted([i.strip() for i in tmp_ips], key=lambda x: int(''.join(
            (lambda a: lambda v: a(a, v))(lambda s, x: x if len(x) == 3 else s(s, '0'+x))(i) for i in x.split('.'))))

    # find the mark
    for table in doc.tables:
        if "ipList" in table.cell(0, 0).paragraphs[0].text:

            # add list table
            main_cell = table.cell(0, 0)
            Parag.delete(main_cell.paragraphs[0])

            cols, ip_cnt = 0, len(tmp_ips)
            lines = 20 if len(tmp_ips) < 4*20 else ceil(len(tmp_ips)/4)

            cols = ceil(ip_cnt / lines)
            rows = ceil(ip_cnt / cols)

            table = main_cell.add_table(1+rows, cols)
            table.style = 'Table Web 1'
            table.allow_autofit = True
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            cell = table.cell(0, 0).merge(table.cell(0, cols-1))
            cell.text = "IP 清單"
            Table.set_cell_color(cell, GRAY)

            for i, ip in enumerate(tmp_ips):
                idx = i // rows
                table.cell(1 + (i % rows), idx).text = ip

            Table.set_content_font(table, size=Pt(
                12), align=1 if ip_cnt < 2*20 else 0)
            Parag.set_font(cell.paragraphs, bold=True, align=1)

            # title = doc.tables[1].cell(0, 0).paragraphs[0]
            # title.text = "表3-1: 主機弱點掃描網段清單"
            # Parag.set_font(title)


# In[8]:


def high_risk_ip_list(doc, wb, sheet, ips=[]):

    tmp_ips = copy.deepcopy(ips)
    empty = True

    # calculate needed data
    if not tmp_ips:
        for row in wb[sheet].values:
            ip, risk = row[0], row[7]
            if ip not in tmp_ips and risk == "高":
                tmp_ips.append(ip)

        if tmp_ips:
            empty = False

            # sort ip
            tmp_ips = sorted([i.strip() for i in tmp_ips], key=lambda x: int(''.join(
                (lambda a: lambda v: a(a, v))(lambda s, x: x if len(x) == 3 else s(s, '0'+x))(i) for i in x.split('.'))))

    else:
        for row in wb[sheet].values:
            ip, risk = row[0], row[7]
            if ip == tmp_ips[0] and risk == "高":
                empty = False

    # find the mark
    for table in doc.tables:
        if "highRiskIpList" in table.cell(0, 0).paragraphs[0].text:
            if not empty:
                # add list table
                main_cell = table.cell(0, 0)
                Parag.delete(main_cell.paragraphs[0])

                table = main_cell.add_table(2, 4)
                table.style = 'Table Web 1'
                cell = table.cell(0, 0).merge(table.cell(0, 3))
                cell.text = "高風險主機 IP 清單"
                Table.set_cell_color(cell, GRAY)

                for i in range(4):
                    Parag.delete(table.cell(1, i).paragraphs[0])

                table.cell(1, 3).add_paragraph(
                    "根據國際標準CVSS 第三版風險評比所述，左列IP屬於高風險設備，基本上容易被入侵。")

                for i, ip in enumerate(sorted(tmp_ips)):
                    idx = i // ceil(len(tmp_ips)/3)
                    table.cell(1, idx).add_paragraph(ip)

                Table.set_content_font(table, align=0)
                Parag.set_font(cell.paragraphs, bold=True)

            else:
                table.cell(0, 0).paragraphs[0].text = "無"

            # title = table.cell(0, 0).add_paragraph()
            # title.text = "表: 主機弱點掃描網段清單"
            # Parag.set_font(title)

    # 替換標定資料
    for parag in doc.paragraphs:
        if "highRiskCount" in parag.text:
            parag.text = parag.text.replace(
                "highRiskCount", str(len(tmp_ips)), 1)


# In[9]:


def cve_describe(doc, wb, sheet, ips=[], risk='高'):

    dic, describe, describe2 = {}, [], []
    level, idx = "", 0
    tmp_ips = copy.deepcopy(ips)

    # calculate needed data
    for row in wb[sheet].values:
        if tmp_ips:
            for ip in tmp_ips:
                if row[0] == ip:
                    level = row[7]
        else:
            level = row[7]

        if level == risk:

            # 如果弱點沒出現過，則記在字典
            if row[8] not in dic:
                dic.setdefault(row[8], idx)
                describe.append([idx+1, "立即的", row[7]])
                describe2.append([row[8]+"\n", row[9], row[10]+"\n", [row[0]]])
                idx += 1

            # 如果弱點出現過，且依據的 IP 沒存在 describe2 中，則新增進去
            elif row[0] not in describe2[dic[row[8]]][3]:
                describe2[dic[row[8]]][3].append(row[0])

        level = ""  # reset

    # sort ip
    for c in range(idx):
        describe2[c][3] = sorted([i.strip() for i in describe2[c][3]], key=lambda x: int(''.join(
            (lambda a: lambda v: a(a, v))(lambda s, x: x if len(x) == 3 else s(s, '0'+x))(i) for i in x.split('.'))))

    # find the mark
    for table in doc.tables:
        if "cveDescribe" in table.cell(0, 0).paragraphs[0].text:
            if idx:
                # add describe table
                main_cell = table.cell(0, 0)
                Parag.delete(main_cell.paragraphs[0])

                for cnt in range(idx):

                    des_table = main_cell.add_table(10, 3)
                    des_table.style = "Table Web 1"
                    des_table.allow_autofit = True
                    des_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    Table.col_widths(des_table, 5.2, 5.2, 5.2)

                    for i in range(7):
                        des_table.cell(i+2, 0).merge(des_table.cell(i+2, 2))

                    tag = ['序號', "急迫性", "風險性"]
                    for i, title in enumerate(tag):
                        cell = des_table.cell(0, i)
                        cell.text = title
                        Table.set_cell_color(cell, GRAY)
                        Parag.set_font(cell.paragraphs, bold=True)

                        cell2 = des_table.cell(1, i)
                        cell2.text = str(describe[cnt][i])
                        Parag.set_font(cell2.paragraphs)

                    tag2 = ["弱點名稱", "弱點說明", "改善建議", "受影響 IP 位置"]
                    for i, title in enumerate(tag2):
                        cell = des_table.cell(2*(i+1), 0)
                        cell.text = title
                        Table.set_cell_color(cell, GRAY)
                        Parag.set_font(cell.paragraphs, bold=True)

                        if i < 3:
                            cell2 = des_table.cell(2*(i+1)+1, 0)
                            cell2.text = str(describe2[cnt][i])
                            Parag.set_font(cell2.paragraphs[0], align=0)
                        else:
                            row_cnt = ceil(len(describe2[cnt][3]) / 3)
                            for j, ip in enumerate(describe2[cnt][3]):
                                idx = j // row_cnt
                                cell2 = des_table.cell(2*(i+1)+1, idx)
                                cell2.paragraphs[0].add_run(
                                    describe2[cnt][3][j])

                                # 若不是每一欄最後一個ip，則加上換行
                                if j % row_cnt + 1 < row_cnt:
                                    cell2.paragraphs[0].add_run("\n")

                            for j in range(0, 3):
                                Parag.set_font(des_table.cell(
                                    9, j).paragraphs, align=WD_ALIGN_PARAGRAPH.LEFT)
            else:
                table.cell(0, 0).paragraphs[0].text = "無"

# In[10]:


def ExcelToWord(excel, sheet, word, company, date, ips, save):

    wb = load_workbook(excel)
    doc = Document(word)

    ip_list(doc, wb, sheet, ips)
    risk_bar(doc, wb, sheet, ips)
    high_risk_ip_list(doc, wb, sheet, ips)
    cve_describe(doc, wb, sheet, ips)

    for parag in doc.paragraphs:
        row = parag.text
        if "YYYY" in row or "MMMM" in row or "DDDD" in row:
            row = row.replace("YYYY", date[0], 1)
            row = row.replace("MMMM", date[1], 1)
            row = row.replace("DDDD", date[2], 1)
            parag.text = row
            Parag.set_font(parag, align=0)

        if "OOOOO" in row:
            tmp_str = row.split("OOOOO")
            parag.clear()
            parag.add_run(tmp_str[0])
            for i in range(1, len(tmp_str)):
                parag.add_run(company).underline = True
                parag.add_run(tmp_str[i])

    doc.save(save)


# %%
